"""
Word Raporu Oluşturma Modülü
Kalp Hastalığı Tahmini Projesi

Bu modül, proje sonuçlarını profesyonel bir Word belgesi olarak oluşturur.

Rapor İçeriği:
- Proje başlığı ve öğrenci bilgileri
- Veri seti özellikleri
- Kullanılan metodoloji
- Model performans tablosu
- Görselleştirmeler (confusion matrix, ROC curves)
- Sonuç ve değerlendirme

Çıktı: Final_Rapor_Heart_Disease.docx

Geliştirici: Ferhat Bayutmuş
Tarih: 17 Kasım 2025
"""

import pandas as pd
import logging
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import os

# Logging ayarları
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('report_generator.log'),
        logging.StreamHandler()
    ]
)

def add_heading_with_style(doc, text, level=1):
    """
    Stilize edilmiş başlık ekler.
    
    Args:
        doc: Document nesnesi
        text (str): Başlık metni
        level (int): Başlık seviyesi
    """
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_style(doc, text, bold=False, italic=False):
    """
    Stilize edilmiş paragraf ekler.
    
    Args:
        doc: Document nesnesi
        text (str): Paragraf metni
        bold (bool): Kalın yazı
        italic (bool): İtalik yazı
    """
    paragraph = doc.add_paragraph(text)
    if bold:
        paragraph.runs[0].font.bold = True
    if italic:
        paragraph.runs[0].font.italic = True
    return paragraph

def add_metrics_table(doc, metrics_df):
    """
    Performans metrikleri tablosunu ekler.
    
    Args:
        doc: Document nesnesi
        metrics_df: Metrik DataFrame'i
    """
    # Tablo oluştur
    table = doc.add_table(rows=1, cols=len(metrics_df.columns) + 1)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Başlık satırı
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Model'
    for i, col in enumerate(metrics_df.columns):
        header_cells[i + 1].text = col.title()
    
    # Başlık satırını kalın yap
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Veri satırları
    for index, row in metrics_df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = index
        for i, value in enumerate(row.values):
            row_cells[i + 1].text = f"{value:.4f}"
    
    return table

def add_image_with_caption(doc, image_path, caption, width=Inches(6)):
    """
    Resmi ve açıklamasını ekler.
    
    Args:
        doc: Document nesnesi
        image_path (str): Resim dosya yolu
        caption (str): Resim açıklaması
        width: Resim genişliği
    """
    try:
        if os.path.exists(image_path):
            # Resmi ekle
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.add_picture(image_path, width=width)
            
            # Açıklama ekle
            caption_paragraph = doc.add_paragraph()
            caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_paragraph.add_run(caption)
            caption_run.font.italic = True
            caption_run.font.size = Pt(10)
            
            logging.info(f"Resim eklendi: {image_path}")
        else:
            doc.add_paragraph(f"⚠ Resim bulunamadı: {image_path}")
            logging.warning(f"Resim bulunamadı: {image_path}")
    except Exception as e:
        logging.error(f"Resim ekleme hatası ({image_path}): {e}")
        doc.add_paragraph(f"❌ Resim yükleme hatası: {os.path.basename(image_path)}")

def generate_report(metrics_df, best_model_name: str, figures_paths: list, 
                   output_path: str = "results/Final_Rapor_Heart_Disease.docx"):
    """
    Kapsamlı proje raporunu oluşturur.
    
    Args:
        metrics_df: Model performans metrikleri
        best_model_name (str): En iyi model adı
        figures_paths (list): Grafik dosya yolları listesi
        output_path (str): Çıktı dosya yolu
    """
    try:
        doc = Document()
        
        # === KAPAK SAYFASI ===
        title = doc.add_heading('Kalp Hastalığı Tahmini', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_heading('Veri Bilimine Giriş Final Projesi', 1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Öğrenci bilgileri
        doc.add_paragraph()
        student_info = doc.add_paragraph()
        student_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        student_run = student_info.add_run("Öğrenci: Ferhat Bayutmuş\nİstanbul Medeniyet Üniversitesi")
        student_run.font.size = Pt(14)
        
        # Tarih
        date_info = doc.add_paragraph()
        date_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_info.add_run(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")
        date_run.font.size = Pt(12)
        
        # Sayfa sonuna git
        doc.add_page_break()
        
        # === ÖZET ===
        add_heading_with_style(doc, 'ÖZET', 1)
        
        ozet_text = f"""Bu projede, UCI Machine Learning Repository'den alınan kalp hastalığı veri seti kullanılarak 
makine öğrenmesi algoritmaları ile kalp hastalığı tahmini yapılmıştır. Proje kapsamında Logistic Regression, 
Random Forest ve K-Nearest Neighbors algoritmaları karşılaştırılmış, en iyi performansı {best_model_name} 
modeli göstermiştir.

Veri seti toplam {len(metrics_df)} farklı algoritma ile analiz edilmiş, cross-validation ve hiperparametre 
optimizasyonu teknikleri kullanılmıştır. Sonuçlar accuracy, precision, recall, F1-score ve AUC metrikleri 
ile değerlendirilmiştir."""
        
        add_paragraph_with_style(doc, ozet_text)
        
        # === VERİ SETİ AÇIKLAMASI ===
        add_heading_with_style(doc, 'VERİ SETİ AÇIKLAMASI', 1)
        
        veri_aciklama = """Bu çalışmada UCI Machine Learning Repository'den alınan "Heart Disease UCI" 
veri seti kullanılmıştır. Veri seti kalp hastalığı teşhisi için gerekli olan çeşitli tıbbi parametreleri içermektedir.

Temel Özellikler:
• Yaş (age): Hastanın yaşı
• Cinsiyet (sex): Hastanın cinsiyeti (1 = erkek, 0 = kadın)
• Göğüs ağrısı tipi (cp): Göğüs ağrısının tipi (0-3 arası)
• Dinlenme kan basıncı (trestbps): Dinlenme halinde kan basıncı
• Kolesterol (chol): Serum kolesterol seviyesi mg/dl
• Açlık kan şekeri (fbs): Açlık kan şekeri > 120 mg/dl (1 = doğru, 0 = yanlış)
• Dinlenme EKG (restecg): Dinlenme elektrokardiogram sonuçları (0-2 arası)
• Maksimum kalp atış hızı (thalach): Ulaşılan maksimum kalp atış hızı
• Egzersiz anjinası (exang): Egzersizle indüklenen anjina (1 = evet, 0 = hayır)
• ST depresyonu (oldpeak): Egzersizle indüklenen ST depresyonu
• ST segment eğimi (slope): Tepe egzersiz ST segmentinin eğimi
• Ana damar sayısı (ca): Floroskopi ile renklendirilmiş ana damar sayısı (0-3)
• Thal: Talassemi (3 = normal; 6 = sabit defekt; 7 = geri döndürülebilir defekt)

Hedef Değişken:
• target: Kalp hastalığı durumu (0 = sağlıklı, 1 = hasta)"""
        
        add_paragraph_with_style(doc, veri_aciklama)
        
        # === VERİ ÖN İŞLEME ===
        add_heading_with_style(doc, 'VERİ ÖN İŞLEME', 1)
        
        on_isleme_text = """Veri analizi sürecinde aşağıdaki ön işleme adımları uygulanmıştır:

1. Eksik Değer Analizi:
   • Veri setindeki her sütun için eksik değer kontrolü yapılmıştır
   • Sayısal değişkenlerdeki eksik değerler median değer ile doldurulmuştur
   • Kategorik değişkenlerdeki eksik değerler mode değeri ile doldurulmuştur

2. Aykırı Değer Tespiti:
   • IQR (Interquartile Range) yöntemi kullanılarak aykırı değerler tespit edilmiştir
   • Aykırı değerler raporlanmış ancak veri setinden çıkarılmamıştır

3. Veri Dönüştürme:
   • Kategorik değişkenler Label Encoding yöntemi ile sayısal değerlere dönüştürülmüştür
   • Tüm özellikler MinMaxScaler kullanılarak 0-1 arasında ölçeklendirilmiştir

4. Veri Bölme:
   • Veri seti %80 eğitim, %20 test olarak bölünmüştür
   • Stratified sampling kullanılarak sınıf dengesinin korunması sağlanmıştır"""
        
        add_paragraph_with_style(doc, on_isleme_text)
        
        # === MODELLER VE SONUÇLAR ===
        add_heading_with_style(doc, 'MODELLER VE SONUÇLAR', 1)
        
        model_text = f"""Bu çalışmada üç farklı makine öğrenmesi algoritması karşılaştırılmıştır:

1. Logistic Regression: Doğrusal sınıflandırma algoritması
2. Random Forest: Ensemble learning algoritması  
3. K-Nearest Neighbors: Mesafe tabanlı sınıflandırma algoritması

Tüm modeller için hiperparametre optimizasyonu GridSearchCV ile yapılmış, 5-fold cross validation 
kullanılmıştır. En iyi performansı {best_model_name} modeli göstermiştir.

Aşağıdaki tabloda tüm modellerin detaylı performans metrikleri yer almaktadır:"""
        
        add_paragraph_with_style(doc, model_text)
        
        # Metrik tablosunu ekle
        add_metrics_table(doc, metrics_df)
        
        # En iyi model vurgusu
        if best_model_name:
            # Model isim eşleştirmesi
            model_name_mapping = {
                'logreg': 'Logistic Regression',
                'rf': 'Random Forest',
                'knn': 'K-Nearest Neighbors'
            }
            display_name = model_name_mapping.get(best_model_name, best_model_name)
            
            # Eğer DataFrame'de display_name varsa kullan
            if display_name in metrics_df.index and 'f1' in metrics_df.columns:
                best_score = metrics_df.loc[display_name, 'f1']
                best_text = f"\n🏆 En İyi Model: {display_name} (F1-Score: {best_score:.4f})"
                add_paragraph_with_style(doc, best_text, bold=True)
            else:
                best_text = f"\n🏆 En İyi Model: {display_name}"
                add_paragraph_with_style(doc, best_text, bold=True)
        
        # === GRAFİKLER ===
        add_heading_with_style(doc, 'GRAFİKLER VE GÖRSELLEŞTİRMELER', 1)
        
        # Grafikleri ekle
        graph_info = {
            'corr_heatmap.png': 'Şekil 1: Özellikler Arası Korelasyon Isı Haritası',
            'class_distribution.png': 'Şekil 2: Hedef Değişken Sınıf Dağılımı',
            'boxplots.png': 'Şekil 3: Sayısal Değişkenlerin Box Plot Analizi',
            'confusion_logistic_regression.png': 'Şekil 4: Logistic Regression Confusion Matrix',
            'confusion_random_forest.png': 'Şekil 5: Random Forest Confusion Matrix',
            'confusion_k-nearest_neighbors.png': 'Şekil 6: K-Nearest Neighbors Confusion Matrix',
            'roc_logistic_regression.png': 'Şekil 7: Logistic Regression ROC Eğrisi',
            'roc_random_forest.png': 'Şekil 8: Random Forest ROC Eğrisi',
            'roc_k-nearest_neighbors.png': 'Şekil 9: K-Nearest Neighbors ROC Eğrisi'
        }
        
        for graph_file, caption in graph_info.items():
            graph_path = os.path.join("results", graph_file)
            if os.path.exists(graph_path):
                add_image_with_caption(doc, graph_path, caption)
                doc.add_paragraph()  # Boşluk için
        
        # === TARTIŞMA VE SONUÇ ===
        add_heading_with_style(doc, 'TARTIŞMA VE SONUÇ', 1)
        
        tartisma_text = f"""Bu çalışmada kalp hastalığı tahmini için üç farklı makine öğrenmesi algoritması 
karşılaştırılmış ve {best_model_name} modeli en iyi performansı göstermiştir.

Model Performansı:
Elde edilen sonuçlar, makine öğrenmesi algoritmalarının kalp hastalığı teşhisinde etkili bir araç 
olarak kullanılabileceğini göstermektedir. Özellikle {best_model_name} modelinin yüksek doğruluk 
oranı, klinik uygulamalarda destek karar sistemi olarak kullanılma potansiyelini ortaya koymaktadır.

Önemli Özellikler:
Korelasyon analizi sonucunda, yaş, maksimum kalp atış hızı, göğüs ağrısı tipi ve ST depresyonu 
gibi değişkenlerin kalp hastalığı tahmini için kritik öneme sahip olduğu görülmüştür.

Klinik Değer:
Bu modelin klinik ortamda kullanılması halinde, doktorların tanı koyma süreçlerini destekleyebilir 
ve erken teşhis oranlarını artırabilir. Ancak, modelin klinik kullanımından önce daha büyük ve 
çeşitli veri setleri ile doğrulanması gerekmektedir.

Gelecek Çalışmalar:
• Daha büyük veri setleri ile model validasyonu
• Deep learning algoritmalarının denenmesi  
• Farklı hastane ve popülasyonlardan veri toplanması
• Real-time klinik entegrasyon çalışmaları"""
        
        add_paragraph_with_style(doc, tartisma_text)
        
        # === KAYNAKÇA ===
        add_heading_with_style(doc, 'KAYNAKÇA', 1)
        
        kaynakca_text = """1. UCI Machine Learning Repository - Heart Disease Dataset
   https://archive.ics.uci.edu/ml/datasets/Heart+Disease

2. Scikit-learn: Machine Learning in Python
   Pedregosa et al., Journal of Machine Learning Research 12, pp. 2825-2830, 2011

3. American Heart Association - Heart Disease and Stroke Statistics
   https://www.heart.org/

4. Pandas: Python Data Analysis Library
   https://pandas.pydata.org/

5. Matplotlib: Python 2D plotting library
   https://matplotlib.org/"""
        
        add_paragraph_with_style(doc, kaynakca_text)
        
        # Belgeyi kaydet
        doc.save(output_path)
        
        logging.info(f"Rapor başarıyla oluşturuldu: {output_path}")
        print(f"📄 Rapor oluşturuldu: {output_path}")
        
        return True
        
    except Exception as e:
        logging.error(f"Rapor oluşturma hatası: {e}")
        print(f"❌ Rapor oluşturma hatası: {e}")
        return False

def create_summary_report(metrics_df, best_model_name):
    """
    Kısa özet rapor oluşturur.
    
    Args:
        metrics_df: Metrik DataFrame'i
        best_model_name (str): En iyi model adı
    """
    print("\n" + "="*60)
    print("📋 PROJE ÖZETİ RAPORU")
    print("="*60)
    print(f"🔬 Veri Seti: Heart Disease UCI Dataset")
    print(f"🤖 Test Edilen Modeller: {len(metrics_df)} adet")
    print(f"🏆 En İyi Model: {best_model_name}")
    
    if best_model_name and best_model_name in metrics_df.index:
        best_metrics = metrics_df.loc[best_model_name]
        print(f"📊 En İyi Model Performansı:")
        print(f"   • Accuracy: {best_metrics['accuracy']:.4f}")
        print(f"   • Precision: {best_metrics['precision']:.4f}")  
        print(f"   • Recall: {best_metrics['recall']:.4f}")
        print(f"   • F1-Score: {best_metrics['f1']:.4f}")
        print(f"   • AUC: {best_metrics['auc']:.4f}")
    
    print("="*60)
    print("✅ Proje başarıyla tamamlandı!")
    print("📁 Tüm dosyalar 'results/' klasöründe kayıtlıdır.")
    print("="*60)

if __name__ == "__main__":
    # Test için örnek kullanım
    print("Rapor oluşturucu modülü test edildi.")
    print("Bu modül main.py tarafından çağrılmalıdır.")